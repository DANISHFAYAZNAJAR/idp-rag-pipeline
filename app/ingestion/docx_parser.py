from pathlib import Path

import structlog
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.ingestion.parsed_document import ParsedDocument, ParsedPage

log = structlog.get_logger()


def _iter_block_items(parent):
    """Yield paragraphs and tables in document body order."""
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _table_to_markdown(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _extract_all_tables(doc: DocxDocument) -> list[str]:
    """Collect top-level and nested table text."""
    seen_ids: set[int] = set()
    tables_md: list[str] = []

    def add_table(table: Table) -> None:
        tid = id(table._tbl)
        if tid in seen_ids:
            return
        seen_ids.add(tid)
        md = _table_to_markdown(table)
        if md.strip():
            tables_md.append(md)
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    add_table(nested)

    for table in doc.tables:
        add_table(table)
    return tables_md


def _is_heading1(paragraph: Paragraph) -> bool:
    style = paragraph.style.name if paragraph.style else ""
    return style == "Heading 1" or style.startswith("Heading 1")


def _pages_have_content(pages: list[ParsedPage]) -> bool:
    return any(p.text.strip() or p.tables for p in pages)


def _parse_docx_blocks(doc: DocxDocument) -> list[ParsedPage]:
    pages: list[ParsedPage] = []
    current_text: list[str] = []
    current_tables: list[str] = []
    page_number = 1

    def flush_page() -> None:
        nonlocal page_number, current_text, current_tables
        text = "\n\n".join(t for t in current_text if t.strip()).strip()
        if text or current_tables:
            pages.append(
                ParsedPage(
                    page_number=page_number,
                    text=text,
                    tables=list(current_tables),
                )
            )
            page_number += 1
        current_text = []
        current_tables = []

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if pages or current_text or current_tables:
                if _is_heading1(block):
                    flush_page()
            current_text.append(text)
        elif isinstance(block, Table):
            md = _table_to_markdown(block)
            if md.strip():
                current_tables.append(md)

    flush_page()
    return pages


def _parse_docx_simple(doc: DocxDocument) -> list[ParsedPage]:
    """Fallback when body-order iteration misses content."""
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = _extract_all_tables(doc)
    if not paragraphs and not tables:
        return []
    return [
        ParsedPage(
            page_number=1,
            text="\n\n".join(paragraphs),
            tables=tables,
        )
    ]


def _parse_docx_unstructured(file_path: Path) -> list[ParsedPage]:
    from unstructured.partition.docx import partition_docx

    elements = partition_docx(filename=str(file_path))
    paragraphs: list[str] = []
    tables: list[str] = []

    for el in elements:
        text = (getattr(el, "text", None) or "").strip()
        if not text:
            continue
        category = getattr(el, "category", "") or type(el).__name__
        if "table" in category.lower():
            tables.append(text)
        else:
            paragraphs.append(text)

    if not paragraphs and not tables:
        return []

    # Split long unstructured output into pseudo-pages (~4000 chars).
    pages: list[ParsedPage] = []
    page_num = 1
    buf: list[str] = []
    buf_len = 0
    for para in paragraphs:
        if buf and buf_len + len(para) > 4000:
            pages.append(ParsedPage(page_number=page_num, text="\n\n".join(buf), tables=[]))
            page_num += 1
            buf, buf_len = [], 0
        buf.append(para)
        buf_len += len(para)
    if buf:
        pages.append(ParsedPage(page_number=page_num, text="\n\n".join(buf), tables=[]))
        page_num += 1

    if tables:
        pages.append(
            ParsedPage(page_number=page_num, text="", tables=tables)
        )

    return pages


def parse_docx(file_path: Path) -> ParsedDocument:
    log.info("docx_parser.start", path=str(file_path))
    doc = DocxDocument(file_path)

    props = doc.core_properties
    metadata = {
        "title": props.title or "",
        "author": props.author or "",
        "creator": props.author or "",
        "total_pages": 0,
    }

    pages = _parse_docx_blocks(doc)
    if not _pages_have_content(pages):
        log.warning("docx_parser.blocks_empty", path=str(file_path))
        pages = _parse_docx_simple(doc)
    if not _pages_have_content(pages):
        log.warning("docx_parser.simple_empty", path=str(file_path))
        try:
            pages = _parse_docx_unstructured(file_path)
        except Exception as exc:
            log.warning("docx_parser.unstructured_failed", error=str(exc))

    if not _pages_have_content(pages):
        raise ValueError(
            "Could not extract any text or tables from this DOCX file. "
            "It may be empty, password-protected, or use an unsupported format."
        )

    metadata["total_pages"] = len(pages)
    log.info(
        "docx_parser.complete",
        total_pages=len(pages),
        text_chars=sum(len(p.text) for p in pages),
        tables=sum(len(p.tables) for p in pages),
    )
    return ParsedDocument(pages=pages, total_pages=len(pages), metadata=metadata)
